from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai

class PDFTranslator:
    def __init__(self, api_key):
        openai.api_key = api_key

    def replace_text_in_run(self, run, old_text, new_text):
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)

    def process_and_translate_text(self, text):
        print(text)
        prompt = """You are an expert translator and you are going to translate a medical document. Please follow these instructions while translating:
            - No need to translate numbers, if numbers are present in the text, skip it. 
            - Do not translate acronyms.
            - Do not translate dates.
            - Do not translate units
            - Maintain the original formatting and structure of the text.
            - If there are no text to translate, don't translate. 
            - Only provide the translation without any additional comments or instructions.
            Translate the following English text to Traditional Chinese:
            
            """
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo-0125",
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": text}]
        )
        translated_text = response.choices[0].message.content.strip()
        print(translated_text)
        return translated_text

    def handle_paragraphs(self, paragraphs, table=False):
        for para in paragraphs:
            current_style = None
            current_text = []
            count = 0

            for run in para.runs:
                if run.style == current_style:
                    current_text.append(run.text)
                else:
                    if current_text:
                        text_to_translate = ''.join(current_text)
                        if text_to_translate.strip():
                            translated_text = self.process_and_translate_text(text_to_translate)
                            self.replace_text_in_run(run, text_to_translate, translated_text)

                    current_style = run.style
                    current_text = [run.text]

                if current_text and (count == len(para.runs) - 1):
                    text_to_translate = ''.join(current_text)
                    if text_to_translate.strip():
                        translated_text = self.process_and_translate_text(text_to_translate)
                        if table:
                            self.replace_text_in_run(run, text_to_translate, translated_text)
                        else:
                            self.replace_text_in_run(para, text_to_translate, translated_text)
                
                count += 1

    def translate_document(self, input_path, output_path):
        doc = Document(input_path)

        self.handle_paragraphs(doc.paragraphs)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.handle_paragraphs(cell.paragraphs, True)

        doc.save(output_path)


